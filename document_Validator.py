import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Inches
import io
import re
from typing import Dict, List, Any, Tuple
from dataclasses import dataclass
from enum import Enum
import json

class SectionType(Enum):
    HEADER = "header"
    FOOTER = "footer"
    PARAGRAPH = "paragraph"
    TABLE = "table"
    LIST = "list"
    IMAGE = "image"
    TOC = "table_of_contents"

@dataclass
class ValidationRule:
    section_type: SectionType
    rule_type: str
    expected_value: Any
    description: str
    is_required: bool = True

class DocumentAnalyzer:
    def __init__(self):
        self.validation_rules = []
    
    def analyze_template(self, doc_file) -> List[ValidationRule]:
        """Analyze template document and extract validation rules"""
        doc = Document(doc_file)
        rules = []
        
        # Analyze header and footer
        rules.extend(self._analyze_header_footer(doc))
        
        # Analyze paragraphs
        rules.extend(self._analyze_paragraphs(doc))
        
        # Analyze tables
        rules.extend(self._analyze_tables(doc))
        
        # Analyze lists
        rules.extend(self._analyze_lists(doc))
        
        # Analyze document structure
        rules.extend(self._analyze_structure(doc))
        
        self.validation_rules = rules
        return rules
    
    def _analyze_header_footer(self, doc) -> List[ValidationRule]:
        """Analyze headers and footers"""
        rules = []
        
        # Check sections for headers/footers
        for section in doc.sections:
            # Header analysis
            if section.header.paragraphs:
                header_text = []
                for para in section.header.paragraphs:
                    if para.text.strip():
                        header_text.append(para.text.strip())
                
                if header_text:
                    rules.append(ValidationRule(
                        section_type=SectionType.HEADER,
                        rule_type="content_structure",
                        expected_value=len(header_text),
                        description=f"Header should contain {len(header_text)} paragraph(s)"
                    ))
            
            # Footer analysis
            if section.footer.paragraphs:
                footer_text = []
                for para in section.footer.paragraphs:
                    if para.text.strip():
                        footer_text.append(para.text.strip())
                
                if footer_text:
                    rules.append(ValidationRule(
                        section_type=SectionType.FOOTER,
                        rule_type="content_structure",
                        expected_value=len(footer_text),
                        description=f"Footer should contain {len(footer_text)} paragraph(s)"
                    ))
        
        return rules
    
    def _analyze_paragraphs(self, doc) -> List[ValidationRule]:
        """Analyze paragraph structure and styles"""
        rules = []
        paragraph_styles = {}
        heading_count = 0
        
        for para in doc.paragraphs:
            if para.text.strip():
                style_name = para.style.name if para.style else "Normal"
                
                if style_name not in paragraph_styles:
                    paragraph_styles[style_name] = 0
                paragraph_styles[style_name] += 1
                
                # Check for headings
                if 'heading' in style_name.lower() or para.style.name.startswith('Heading'):
                    heading_count += 1
        
        # Add rules for paragraph styles
        for style, count in paragraph_styles.items():
            rules.append(ValidationRule(
                section_type=SectionType.PARAGRAPH,
                rule_type="style_usage",
                expected_value={"style": style, "min_count": 1},
                description=f"Document should use '{style}' style",
                is_required=count > 2  # Only require styles used more than twice
            ))
        
        if heading_count > 0:
            rules.append(ValidationRule(
                section_type=SectionType.PARAGRAPH,
                rule_type="heading_structure",
                expected_value=heading_count,
                description=f"Document should have structured headings (found {heading_count} in template)"
            ))
        
        return rules
    
    def _analyze_tables(self, doc) -> List[ValidationRule]:
        """Analyze table structures"""
        rules = []
        table_structures = []
        
        for i, table in enumerate(doc.tables):
            rows = len(table.rows)
            cols = len(table.columns) if table.rows else 0
            
            # Analyze first row for headers
            has_header = False
            header_texts = []
            if table.rows:
                first_row = table.rows[0]
                for cell in first_row.cells:
                    cell_text = cell.text.strip()
                    header_texts.append(cell_text)
                    # Check if cell might be a header (bold, different style, etc.)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.bold:
                                has_header = True
                                break
            
            table_structure = {
                "table_index": i,
                "rows": rows,
                "columns": cols,
                "has_header": has_header,
                "header_texts": header_texts
            }
            table_structures.append(table_structure)
            
            rules.append(ValidationRule(
                section_type=SectionType.TABLE,
                rule_type="table_structure",
                expected_value=table_structure,
                description=f"Table {i+1}: {rows} rows × {cols} columns" + 
                           (" with headers" if has_header else "")
            ))
        
        if table_structures:
            rules.append(ValidationRule(
                section_type=SectionType.TABLE,
                rule_type="table_count",
                expected_value=len(table_structures),
                description=f"Document should contain {len(table_structures)} table(s)"
            ))
        
        return rules
    
    def _analyze_lists(self, doc) -> List[ValidationRule]:
        """Analyze list structures"""
        rules = []
        list_count = 0
        
        for para in doc.paragraphs:
            if para.style and ('list' in para.style.name.lower() or 
                             para.text.strip().startswith(('•', '-', '*', '1.', '2.', '3.'))):
                list_count += 1
        
        if list_count > 0:
            rules.append(ValidationRule(
                section_type=SectionType.LIST,
                rule_type="list_usage",
                expected_value=list_count,
                description=f"Document uses lists (found {list_count} list items in template)",
                is_required=False
            ))
        
        return rules
    
    def _analyze_structure(self, doc) -> List[ValidationRule]:
        """Analyze overall document structure"""
        rules = []
        
        # Count total paragraphs with content
        content_paragraphs = sum(1 for para in doc.paragraphs if para.text.strip())
        
        rules.append(ValidationRule(
            section_type=SectionType.PARAGRAPH,
            rule_type="content_volume",
            expected_value={"min_paragraphs": max(1, content_paragraphs // 2)},
            description=f"Document should have substantial content (minimum {max(1, content_paragraphs // 2)} paragraphs)",
            is_required=False
        ))
        
        return rules
    
    def validate_document(self, doc_file, template_rules: List[ValidationRule]) -> Dict[str, Any]:
        """Validate document against template rules"""
        doc = Document(doc_file)
        results = {
            "total_rules": len(template_rules),
            "passed": 0,
            "failed": 0,
            "warnings": 0,
            "details": []
        }
        
        for rule in template_rules:
            validation_result = self._validate_rule(doc, rule)
            results["details"].append(validation_result)
            
            if validation_result["status"] == "PASS":
                results["passed"] += 1
            elif validation_result["status"] == "FAIL":
                results["failed"] += 1
            else:
                results["warnings"] += 1
        
        return results
    
    def _validate_rule(self, doc, rule: ValidationRule) -> Dict[str, Any]:
        """Validate a single rule against the document"""
        result = {
            "rule_description": rule.description,
            "rule_type": rule.rule_type,
            "section_type": rule.section_type.value,
            "is_required": rule.is_required,
            "status": "PASS",
            "message": "",
            "expected": rule.expected_value,
            "actual": None
        }
        
        try:
            if rule.section_type == SectionType.HEADER:
                result = self._validate_header(doc, rule, result)
            elif rule.section_type == SectionType.FOOTER:
                result = self._validate_footer(doc, rule, result)
            elif rule.section_type == SectionType.PARAGRAPH:
                result = self._validate_paragraph(doc, rule, result)
            elif rule.section_type == SectionType.TABLE:
                result = self._validate_table(doc, rule, result)
            elif rule.section_type == SectionType.LIST:
                result = self._validate_list(doc, rule, result)
            
        except Exception as e:
            result["status"] = "ERROR"
            result["message"] = f"Error during validation: {str(e)}"
        
        return result
    
    def _validate_header(self, doc, rule, result):
        """Validate header rules"""
        header_paragraphs = 0
        for section in doc.sections:
            for para in section.header.paragraphs:
                if para.text.strip():
                    header_paragraphs += 1
        
        result["actual"] = header_paragraphs
        
        if rule.rule_type == "content_structure":
            expected_count = rule.expected_value
            if header_paragraphs != expected_count:
                result["status"] = "FAIL" if rule.is_required else "WARNING"
                result["message"] = f"Expected {expected_count} header paragraphs, found {header_paragraphs}"
            else:
                result["message"] = "Header structure matches template"
        
        return result
    
    def _validate_footer(self, doc, rule, result):
        """Validate footer rules"""
        footer_paragraphs = 0
        for section in doc.sections:
            for para in section.footer.paragraphs:
                if para.text.strip():
                    footer_paragraphs += 1
        
        result["actual"] = footer_paragraphs
        
        if rule.rule_type == "content_structure":
            expected_count = rule.expected_value
            if footer_paragraphs != expected_count:
                result["status"] = "FAIL" if rule.is_required else "WARNING"
                result["message"] = f"Expected {expected_count} footer paragraphs, found {footer_paragraphs}"
            else:
                result["message"] = "Footer structure matches template"
        
        return result
    
    def _validate_paragraph(self, doc, rule, result):
        """Validate paragraph rules"""
        if rule.rule_type == "style_usage":
            expected_style = rule.expected_value["style"]
            style_count = sum(1 for para in doc.paragraphs 
                            if para.text.strip() and para.style.name == expected_style)
            
            result["actual"] = style_count
            min_count = rule.expected_value.get("min_count", 1)
            
            if style_count < min_count:
                result["status"] = "FAIL" if rule.is_required else "WARNING"
                result["message"] = f"Style '{expected_style}' used {style_count} times, expected at least {min_count}"
            else:
                result["message"] = f"Style '{expected_style}' properly used ({style_count} times)"
        
        elif rule.rule_type == "heading_structure":
            heading_count = sum(1 for para in doc.paragraphs 
                              if para.text.strip() and 'heading' in para.style.name.lower())
            result["actual"] = heading_count
            
            if heading_count == 0:
                result["status"] = "FAIL" if rule.is_required else "WARNING"
                result["message"] = "No headings found in document"
            else:
                result["message"] = f"Document has {heading_count} headings"
        
        elif rule.rule_type == "content_volume":
            content_paragraphs = sum(1 for para in doc.paragraphs if para.text.strip())
            result["actual"] = content_paragraphs
            min_paragraphs = rule.expected_value["min_paragraphs"]
            
            if content_paragraphs < min_paragraphs:
                result["status"] = "WARNING"
                result["message"] = f"Document has {content_paragraphs} paragraphs, recommended minimum {min_paragraphs}"
            else:
                result["message"] = f"Document has adequate content ({content_paragraphs} paragraphs)"
        
        return result
    
    def _validate_table(self, doc, rule, result):
        """Validate table rules"""
        if rule.rule_type == "table_count":
            actual_count = len(doc.tables)
            expected_count = rule.expected_value
            result["actual"] = actual_count
            
            if actual_count != expected_count:
                result["status"] = "FAIL" if rule.is_required else "WARNING"
                result["message"] = f"Expected {expected_count} tables, found {actual_count}"
            else:
                result["message"] = f"Correct number of tables ({actual_count})"
        
        elif rule.rule_type == "table_structure":
            table_info = rule.expected_value
            table_index = table_info["table_index"]
            
            if table_index < len(doc.tables):
                table = doc.tables[table_index]
                actual_rows = len(table.rows)
                actual_cols = len(table.columns) if table.rows else 0
                
                result["actual"] = {"rows": actual_rows, "columns": actual_cols}
                
                expected_rows = table_info["rows"]
                expected_cols = table_info["columns"]
                
                if actual_rows != expected_rows or actual_cols != expected_cols:
                    result["status"] = "FAIL" if rule.is_required else "WARNING"
                    result["message"] = f"Table {table_index + 1}: Expected {expected_rows}×{expected_cols}, found {actual_rows}×{actual_cols}"
                else:
                    result["message"] = f"Table {table_index + 1} structure matches template"
            else:
                result["status"] = "FAIL"
                result["message"] = f"Table {table_index + 1} not found in document"
                result["actual"] = "Missing"
        
        return result
    
    def _validate_list(self, doc, rule, result):
        """Validate list rules"""
        list_count = sum(1 for para in doc.paragraphs 
                        if para.style and ('list' in para.style.name.lower() or 
                           para.text.strip().startswith(('•', '-', '*', '1.', '2.', '3.'))))
        
        result["actual"] = list_count
        
        if rule.rule_type == "list_usage":
            if list_count == 0:
                result["status"] = "WARNING"
                result["message"] = "No lists found in document (template contained lists)"
            else:
                result["message"] = f"Document contains {list_count} list items"
        
        return result

def main():
    st.set_page_config(
        page_title="Word Document Analyzer & Validator",
        page_icon="📄",
        layout="wide"
    )
    
    st.title("📄 Word Document Template Analyzer & Validator")
    st.markdown("---")
    
    # Initialize session state
    if 'analyzer' not in st.session_state:
        st.session_state.analyzer = DocumentAnalyzer()
    if 'template_rules' not in st.session_state:
        st.session_state.template_rules = []
    if 'template_analyzed' not in st.session_state:
        st.session_state.template_analyzed = False
    
    # Create two columns for the main interface
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.header("📋 Step 1: Upload Template Document")
        st.markdown("Upload a Word document template to analyze its structure and create validation rules.")
        
        template_file = st.file_uploader(
            "Choose template document",
            type=['docx'],
            key="template_upload"
        )
        
        if template_file is not None:
            if st.button("🔍 Analyze Template", type="primary"):
                with st.spinner("Analyzing template document..."):
                    try:
                        rules = st.session_state.analyzer.analyze_template(template_file)
                        st.session_state.template_rules = rules
                        st.session_state.template_analyzed = True
                        st.success(f"✅ Template analyzed! Found {len(rules)} validation rules.")
                    except Exception as e:
                        st.error(f"❌ Error analyzing template: {str(e)}")
        
        # Display template analysis results
        if st.session_state.template_analyzed and st.session_state.template_rules:
            st.subheader("📊 Template Analysis Results")
            
            # Summary metrics
            total_rules = len(st.session_state.template_rules)
            required_rules = sum(1 for rule in st.session_state.template_rules if rule.is_required)
            optional_rules = total_rules - required_rules
            
            metric_col1, metric_col2, metric_col3 = st.columns(3)
            with metric_col1:
                st.metric("Total Rules", total_rules)
            with metric_col2:
                st.metric("Required", required_rules)
            with metric_col3:
                st.metric("Optional", optional_rules)
            
            # Rules breakdown
            with st.expander("📋 View All Validation Rules"):
                rules_data = []
                for i, rule in enumerate(st.session_state.template_rules):
                    rules_data.append({
                        "Rule #": i + 1,
                        "Section": rule.section_type.value.title(),
                        "Type": rule.rule_type.replace("_", " ").title(),
                        "Description": rule.description,
                        "Required": "Yes" if rule.is_required else "No"
                    })
                
                df_rules = pd.DataFrame(rules_data)
                st.dataframe(df_rules, use_container_width=True, hide_index=True)
    
    with col2:
        st.header("📝 Step 2: Validate Document")
        st.markdown("Upload a Word document to validate against the template rules.")
        
        if not st.session_state.template_analyzed:
            st.info("👆 Please analyze a template document first.")
            st.stop()
        
        document_file = st.file_uploader(
            "Choose document to validate",
            type=['docx'],
            key="document_upload"
        )
        
        if document_file is not None:
            if st.button("✅ Validate Document", type="primary"):
                with st.spinner("Validating document against template rules..."):
                    try:
                        validation_results = st.session_state.analyzer.validate_document(
                            document_file, st.session_state.template_rules
                        )
                        
                        # Display validation summary
                        st.subheader("🎯 Validation Results")
                        
                        # Summary metrics
                        total = validation_results["total_rules"]
                        passed = validation_results["passed"]
                        failed = validation_results["failed"]
                        warnings = validation_results["warnings"]
                        
                        # Calculate success rate
                        success_rate = (passed / total * 100) if total > 0 else 0
                        
                        # Display metrics
                        result_col1, result_col2, result_col3, result_col4 = st.columns(4)
                        with result_col1:
                            st.metric("Passed", passed, delta=f"{success_rate:.1f}%")
                        with result_col2:
                            st.metric("Failed", failed)
                        with result_col3:
                            st.metric("Warnings", warnings)
                        with result_col4:
                            st.metric("Total Rules", total)
                        
                        # Overall status
                        if failed == 0:
                            if warnings == 0:
                                st.success("🎉 Perfect! Document fully complies with template.")
                            else:
                                st.warning(f"✨ Good! Document passes all required rules ({warnings} warnings).")
                        else:
                            st.error(f"❌ Document validation failed ({failed} critical issues).")
                        
                        # Detailed results
                        st.subheader("📋 Detailed Validation Results")
                        
                        # Filter options
                        filter_col1, filter_col2 = st.columns(2)
                        with filter_col1:
                            status_filter = st.selectbox(
                                "Filter by status:",
                                ["All", "PASS", "FAIL", "WARNING"],
                                key="status_filter"
                            )
                        with filter_col2:
                            section_types = list(set(detail["section_type"] for detail in validation_results["details"]))
                            section_filter = st.selectbox(
                                "Filter by section:",
                                ["All"] + section_types,
                                key="section_filter"
                            )
                        
                        # Filter results
                        filtered_details = validation_results["details"]
                        if status_filter != "All":
                            filtered_details = [d for d in filtered_details if d["status"] == status_filter]
                        if section_filter != "All":
                            filtered_details = [d for d in filtered_details if d["section_type"] == section_filter]
                        
                        # Display results
                        for detail in filtered_details:
                            status = detail["status"]
                            
                            # Choose icon and color based on status
                            if status == "PASS":
                                icon = "✅"
                                color = "green"
                            elif status == "FAIL":
                                icon = "❌"
                                color = "red"
                            elif status == "WARNING":
                                icon = "⚠️"
                                color = "orange"
                            else:
                                icon = "❓"
                                color = "gray"
                            
                            with st.container():
                                st.markdown(f"""
                                <div style="padding: 10px; border-left: 4px solid {color}; margin: 5px 0; background-color: rgba(255,255,255,0.1);">
                                    <strong>{icon} {detail['rule_description']}</strong><br>
                                    <small>Section: {detail['section_type'].title()} | Type: {detail['rule_type'].replace('_', ' ').title()}</small><br>
                                    {detail['message']}
                                </div>
                                """, unsafe_allow_html=True)
                        
                        # Export results option
                        if st.button("📊 Export Results as JSON"):
                            json_results = json.dumps(validation_results, indent=2, default=str)
                            st.download_button(
                                label="📥 Download Validation Report",
                                data=json_results,
                                file_name=f"validation_report_{pd.Timestamp.now().strftime('%Y%m%d_%H%M%S')}.json",
                                mime="application/json"
                            )
                    
                    except Exception as e:
                        st.error(f"❌ Error during validation: {str(e)}")
    
    # Footer
    st.markdown("---")
    st.markdown(
        "<div style='text-align: center; color: gray;'>"
        "📄 Word Document Template Analyzer & Validator | "
        "Analyze templates, validate documents, ensure consistency"
        "</div>",
        unsafe_allow_html=True
    )

if __name__ == "__main__":
    main()
