import streamlit as st
import json
import re
from dataclasses import dataclass, asdict
from typing import List, Dict, Any, Optional
from docx import Document
import io
import google.generativeai as genai
from datetime import datetime

# Configure page
st.set_page_config(
    page_title="Document Compiler Agent",
    page_icon="📄",
    layout="wide"
)

@dataclass
class ValidationRule:
    rule_type: str  # 'structural', 'content', 'formatting'
    category: str   # 'presence', 'order', 'hierarchy', 'length', 'keywords'
    section: str
    requirement: str
    severity: str   # 'critical', 'warning', 'suggestion'
    details: Dict[str, Any]

@dataclass
class DocumentSection:
    heading: str
    level: int
    content: str
    word_count: int
    position: int

@dataclass
class ValidationResult:
    rule: ValidationRule
    passed: bool
    message: str
    details: Optional[str] = None

class DocumentParser:
    """Handles parsing of Word documents into structured format"""
    
    def parse_docx(self, uploaded_file) -> List[DocumentSection]:
        """Parse uploaded Word document and extract sections"""
        try:
            doc = Document(uploaded_file)
            sections = []
            current_position = 0
            
            for paragraph in doc.paragraphs:
                if paragraph.style.name.startswith('Heading'):
                    # Extract heading level
                    level = int(paragraph.style.name.split()[-1]) if paragraph.style.name.split()[-1].isdigit() else 1
                    
                    # Get content until next heading
                    content = self._get_section_content(doc, paragraph)
                    word_count = len(content.split()) if content else 0
                    
                    section = DocumentSection(
                        heading=paragraph.text.strip(),
                        level=level,
                        content=content,
                        word_count=word_count,
                        position=current_position
                    )
                    sections.append(section)
                    current_position += 1
            
            return sections
        except Exception as e:
            st.error(f"Error parsing document: {str(e)}")
            return []
    
    def _get_section_content(self, doc, heading_paragraph) -> str:
        """Extract content between current heading and next heading"""
        content_parts = []
        start_collecting = False
        heading_level = self._get_heading_level(heading_paragraph.style.name)
        
        for paragraph in doc.paragraphs:
            if paragraph == heading_paragraph:
                start_collecting = True
                continue
            
            if start_collecting:
                if paragraph.style.name.startswith('Heading'):
                    current_level = self._get_heading_level(paragraph.style.name)
                    if current_level <= heading_level:
                        break
                content_parts.append(paragraph.text)
        
        return '\n'.join(content_parts).strip()
    
    def _get_heading_level(self, style_name) -> int:
        """Extract heading level from style name"""
        if 'Heading' in style_name:
            parts = style_name.split()
            if len(parts) > 1 and parts[-1].isdigit():
                return int(parts[-1])
        return 1

class RuleEngine:
    """Handles validation rules and document checking"""
    
    def __init__(self):
        self.template_structure = self._load_default_template()
        self.validation_rules = self._load_default_rules()
    
    def _load_default_template(self) -> Dict[str, Any]:
        """Load default template structure"""
        return {
            "required_sections": [
                {"heading": "Executive Summary", "level": 1, "min_words": 100},
                {"heading": "Introduction", "level": 1, "min_words": 150},
                {"heading": "Methodology", "level": 1, "min_words": 200},
                {"heading": "Results", "level": 1, "min_words": 300},
                {"heading": "Discussion", "level": 1, "min_words": 200},
                {"heading": "Conclusion", "level": 1, "min_words": 100}
            ],
            "optional_sections": [
                {"heading": "Appendix", "level": 1},
                {"heading": "References", "level": 1}
            ],
            "section_order": [
                "Executive Summary", "Introduction", "Methodology", 
                "Results", "Discussion", "Conclusion", "References", "Appendix"
            ]
        }
    
    def _load_default_rules(self) -> List[ValidationRule]:
        """Load default validation rules"""
        rules = []
        
        # Structural rules
        for section in self.template_structure["required_sections"]:
            rules.append(ValidationRule(
                rule_type="structural",
                category="presence",
                section=section["heading"],
                requirement=f"Section '{section['heading']}' must be present",
                severity="critical",
                details={"required": True}
            ))
            
            if "min_words" in section:
                rules.append(ValidationRule(
                    rule_type="content",
                    category="length",
                    section=section["heading"],
                    requirement=f"Minimum {section['min_words']} words required",
                    severity="warning",
                    details={"min_words": section["min_words"]}
                ))
        
        # Content rules
        content_rules = [
            {
                "section": "Introduction",
                "keywords": ["objective", "purpose", "scope"],
                "requirement": "Must mention project objectives/purpose"
            },
            {
                "section": "Methodology",
                "keywords": ["method", "approach", "procedure"],
                "requirement": "Must describe methodology or approach"
            },
            {
                "section": "Results",
                "keywords": ["result", "finding", "data", "figure", "table"],
                "requirement": "Must present results or findings"
            }
        ]
        
        for rule in content_rules:
            rules.append(ValidationRule(
                rule_type="content",
                category="keywords",
                section=rule["section"],
                requirement=rule["requirement"],
                severity="warning",
                details={"keywords": rule["keywords"]}
            ))
        
        return rules
    
    def validate_document(self, sections: List[DocumentSection]) -> List[ValidationResult]:
        """Validate document against all rules"""
        results = []
        
        for rule in self.validation_rules:
            result = self._apply_rule(rule, sections)
            results.append(result)
        
        return results
    
    def _apply_rule(self, rule: ValidationRule, sections: List[DocumentSection]) -> ValidationResult:
        """Apply a single validation rule"""
        section_dict = {s.heading: s for s in sections}
        
        if rule.category == "presence":
            return self._check_presence(rule, section_dict)
        elif rule.category == "length":
            return self._check_length(rule, section_dict)
        elif rule.category == "keywords":
            return self._check_keywords(rule, section_dict)
        else:
            return ValidationResult(
                rule=rule,
                passed=True,
                message="Rule type not implemented"
            )
    
    def _check_presence(self, rule: ValidationRule, section_dict: Dict[str, DocumentSection]) -> ValidationResult:
        """Check if required section is present"""
        section_present = rule.section in section_dict
        
        return ValidationResult(
            rule=rule,
            passed=section_present,
            message=f"✅ Section '{rule.section}' found" if section_present else f"❌ Missing required section: '{rule.section}'"
        )
    
    def _check_length(self, rule: ValidationRule, section_dict: Dict[str, DocumentSection]) -> ValidationResult:
        """Check if section meets minimum word count"""
        if rule.section not in section_dict:
            return ValidationResult(
                rule=rule,
                passed=False,
                message=f"❌ Cannot check length: Section '{rule.section}' not found"
            )
        
        section = section_dict[rule.section]
        min_words = rule.details.get("min_words", 0)
        meets_requirement = section.word_count >= min_words
        
        return ValidationResult(
            rule=rule,
            passed=meets_requirement,
            message=f"✅ Section '{rule.section}' has {section.word_count} words (≥{min_words})" if meets_requirement 
                   else f"⚠️ Section '{rule.section}' has only {section.word_count} words (minimum: {min_words})"
        )
    
    def _check_keywords(self, rule: ValidationRule, section_dict: Dict[str, DocumentSection]) -> ValidationResult:
        """Check if section contains required keywords"""
        if rule.section not in section_dict:
            return ValidationResult(
                rule=rule,
                passed=False,
                message=f"❌ Cannot check keywords: Section '{rule.section}' not found"
            )
        
        section = section_dict[rule.section]
        keywords = rule.details.get("keywords", [])
        content_lower = section.content.lower()
        
        found_keywords = [kw for kw in keywords if kw.lower() in content_lower]
        has_keywords = len(found_keywords) > 0
        
        return ValidationResult(
            rule=rule,
            passed=has_keywords,
            message=f"✅ Found keywords in '{rule.section}': {', '.join(found_keywords)}" if has_keywords
                   else f"⚠️ No required keywords found in '{rule.section}'. Expected: {', '.join(keywords)}"
        )

class LLMAnalyzer:
    """Handles LLM-based content analysis"""
    
    def __init__(self, api_key: str = None):
        self.api_key = api_key
        if api_key:
            genai.configure(api_key=api_key)
            self.model = genai.GenerativeModel('gemini-pro')
        else:
            self.model = None
    
    def analyze_section_quality(self, section: DocumentSection, expected_purpose: str) -> str:
        """Analyze section quality using LLM"""
        if not self.model:
            return "LLM analysis not available (API key not provided)"
        
        prompt = f"""
        Analyze the following document section for quality and completeness:
        
        Section: {section.heading}
        Expected Purpose: {expected_purpose}
        Word Count: {section.word_count}
        
        Content:
        {section.content}
        
        Please evaluate:
        1. Does this section fulfill its expected purpose?
        2. Is the content clear and well-structured?
        3. What information might be missing?
        4. Any suggestions for improvement?
        
        Provide a concise analysis in 2-3 sentences.
        """
        
        try:
            response = self.model.generate_content(prompt)
            return response.text
        except Exception as e:
            return f"LLM analysis failed: {str(e)}"
    
    def check_document_coherence(self, sections: List[DocumentSection]) -> str:
        """Check overall document coherence"""
        if not self.model:
            return "LLM analysis not available"
        
        # Create a summary of all sections
        section_summaries = []
        for section in sections:
            summary = f"{section.heading}: {section.content[:200]}..." if len(section.content) > 200 else f"{section.heading}: {section.content}"
            section_summaries.append(summary)
        
        prompt = f"""
        Analyze the following document sections for overall coherence and logical flow:
        
        {chr(10).join(section_summaries)}
        
        Please evaluate:
        1. Do the sections flow logically from one to another?
        2. Is there consistency in terminology and approach?
        3. Are there any gaps in the logical progression?
        4. Any recommendations for improving document structure?
        
        Provide a brief assessment in 2-3 sentences.
        """
        
        try:
            response = self.model.generate_content(prompt)
            return response.text
        except Exception as e:
            return f"Coherence analysis failed: {str(e)}"

def main():
    st.title("📄 Document Compiler Agent")
    st.markdown("Upload your Word document to validate it against template requirements and get intelligent feedback.")
    
    # Sidebar for configuration
    with st.sidebar:
        st.header("⚙️ Configuration")
        
        # API Key input
        api_key = st.text_input(
            "Google Gemini API Key (optional)",
            type="password",
            help="Enter your API key for advanced LLM analysis"
        )
        
        # Template configuration
        st.subheader("Template Settings")
        use_custom_template = st.checkbox("Use custom template", help="Upload your own template structure")
        
        if use_custom_template:
            st.info("Custom template functionality would be implemented here")
    
    # Initialize components
    parser = DocumentParser()
    rule_engine = RuleEngine()
    llm_analyzer = LLMAnalyzer(api_key) if api_key else LLMAnalyzer()
    
    # Main interface
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.subheader("📤 Upload Document")
        uploaded_file = st.file_uploader(
            "Choose a Word document (.docx)",
            type=['docx'],
            help="Upload the document you want to validate"
        )
        
        if uploaded_file is not None:
            st.success(f"File uploaded: {uploaded_file.name}")
            
            # Parse document
            with st.spinner("Parsing document..."):
                sections = parser.parse_docx(uploaded_file)
            
            if sections:
                st.info(f"Found {len(sections)} sections in the document")
                
                # Display document structure
                with st.expander("📋 Document Structure"):
                    for i, section in enumerate(sections, 1):
                        st.write(f"{i}. **{section.heading}** (Level {section.level}) - {section.word_count} words")
    
    with col2:
        st.subheader("📊 Template Requirements")
        
        # Display template structure
        template = rule_engine.template_structure
        
        st.write("**Required Sections:**")
        for section in template["required_sections"]:
            min_words = section.get("min_words", "No minimum")
            st.write(f"• {section['heading']} (Min: {min_words} words)")
        
        st.write("**Optional Sections:**")
        for section in template["optional_sections"]:
            st.write(f"• {section['heading']}")
    
    # Validation Results
    if uploaded_file is not None and sections:
        st.header("🔍 Validation Results")
        
        # Run validation
        with st.spinner("Running validation..."):
            validation_results = rule_engine.validate_document(sections)
        
        # Categorize results
        critical_issues = [r for r in validation_results if r.rule.severity == "critical" and not r.passed]
        warnings = [r for r in validation_results if r.rule.severity == "warning" and not r.passed]
        passed_checks = [r for r in validation_results if r.passed]
        
        # Display metrics
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Critical Issues", len(critical_issues))
        with col2:
            st.metric("Warnings", len(warnings))
        with col3:
            st.metric("Passed Checks", len(passed_checks))
        with col4:
            pass_rate = len(passed_checks) / len(validation_results) * 100 if validation_results else 0
            st.metric("Pass Rate", f"{pass_rate:.1f}%")
        
        # Display detailed results
        if critical_issues:
            st.error("**Critical Issues Found:**")
            for result in critical_issues:
                st.write(f"• {result.message}")
        
        if warnings:
            st.warning("**Warnings:**")
            for result in warnings:
                st.write(f"• {result.message}")
        
        if passed_checks:
            with st.expander("✅ Passed Checks"):
                for result in passed_checks:
                    st.write(f"• {result.message}")
        
        # LLM Analysis
        if api_key:
            st.header("🤖 AI-Powered Analysis")
            
            tab1, tab2 = st.tabs(["Section Analysis", "Document Coherence"])
            
            with tab1:
                st.subheader("Section Quality Analysis")
                
                # Define expected purposes for sections
                expected_purposes = {
                    "Executive Summary": "Provide a concise overview of the entire document",
                    "Introduction": "Introduce the topic, objectives, and scope",
                    "Methodology": "Describe the methods and approaches used",
                    "Results": "Present findings, data, and outcomes",
                    "Discussion": "Interpret results and discuss implications",
                    "Conclusion": "Summarize key findings and conclusions"
                }
                
                for section in sections:
                    if section.heading in expected_purposes:
                        with st.expander(f"Analysis: {section.heading}"):
                            with st.spinner(f"Analyzing {section.heading}..."):
                                analysis = llm_analyzer.analyze_section_quality(
                                    section, 
                                    expected_purposes[section.heading]
                                )
                            st.write(analysis)
            
            with tab2:
                st.subheader("Document Coherence Analysis")
                with st.spinner("Analyzing document coherence..."):
                    coherence_analysis = llm_analyzer.check_document_coherence(sections)
                st.write(coherence_analysis)
        
        # Generate Report
        st.header("📄 Validation Report")
        
        if st.button("Generate Detailed Report"):
            report = generate_report(uploaded_file.name, validation_results, sections)
            st.download_button(
                label="Download Report",
                data=report,
                file_name=f"validation_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                mime="text/plain"
            )

def generate_report(filename: str, results: List[ValidationResult], sections: List[DocumentSection]) -> str:
    """Generate a detailed validation report"""
    report_lines = [
        "DOCUMENT VALIDATION REPORT",
        "=" * 50,
        f"Document: {filename}",
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "",
        "DOCUMENT STRUCTURE:",
        "-" * 20
    ]
    
    for i, section in enumerate(sections, 1):
        report_lines.append(f"{i}. {section.heading} (Level {section.level}) - {section.word_count} words")
    
    report_lines.extend([
        "",
        "VALIDATION RESULTS:",
        "-" * 20
    ])
    
    critical_issues = [r for r in results if r.rule.severity == "critical" and not r.passed]
    warnings = [r for r in results if r.rule.severity == "warning" and not r.passed]
    passed_checks = [r for r in results if r.passed]
    
    if critical_issues:
        report_lines.append("\nCRITICAL ISSUES:")
        for result in critical_issues:
            report_lines.append(f"• {result.message}")
    
    if warnings:
        report_lines.append("\nWARNINGS:")
        for result in warnings:
            report_lines.append(f"• {result.message}")
    
    report_lines.append(f"\nSUMMARY:")
    report_lines.append(f"• Total checks: {len(results)}")
    report_lines.append(f"• Passed: {len(passed_checks)}")
    report_lines.append(f"• Warnings: {len(warnings)}")
    report_lines.append(f"• Critical issues: {len(critical_issues)}")
    
    pass_rate = len(passed_checks) / len(results) * 100 if results else 0
    report_lines.append(f"• Pass rate: {pass_rate:.1f}%")
    
    return "\n".join(report_lines)

if __name__ == "__main__":
    main()