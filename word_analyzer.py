import streamlit as st
from docx import Document
from docx.shared import Pt, Cm # Inches can also be used
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
import io
import os

# --- 1. Document Analyzer ---
def get_paragraph_formatting_details(para):
    """Extracts formatting details from a paragraph object."""
    details = {}
    if para.style and para.style.name:
        details['style_name'] = para.style.name

    # Font details from the first run (dominant formatting)
    if para.runs:
        run = para.runs[0]
        font = run.font
        details['font_name'] = font.name
        details['font_size_pt'] = round(font.size.pt,1) if font.size and font.size.pt is not None else None
        details['font_bold'] = font.bold
        details['font_italic'] = font.italic
        details['font_underline'] = font.underline
        details['font_color_rgb'] = str(font.color.rgb) if font.color and font.color.rgb else None

    # Paragraph formatting
    p_format = para.paragraph_format
    details['alignment'] = WD_ALIGN_PARAGRAPH.to_xml(p_format.alignment).lower() if p_format.alignment is not None else 'left' # Default to left if None
    details['space_before_pt'] = round(p_format.space_before.pt,1) if p_format.space_before and p_format.space_before.pt is not None else None
    details['space_after_pt'] = round(p_format.space_after.pt,1) if p_format.space_after and p_format.space_after.pt is not None else None
    details['line_spacing_val'] = round(p_format.line_spacing,2) if p_format.line_spacing is not None else None # This can be a number (e.g., 1.5) or a WD_LINE_SPACING enum value
    details['left_indent_cm'] = round(p_format.left_indent.cm,2) if p_format.left_indent and p_format.left_indent.cm is not None else None
    details['first_line_indent_cm'] = round(p_format.first_line_indent.cm,2) if p_format.first_line_indent and p_format.first_line_indent.cm is not None else None

    return details

def analyze_header_footer_elements(elements_collection):
    """Analyzes paragraphs and tables within a header/footer collection (e.g., section.header)."""
    analysis = {"paragraphs_summary": [], "tables_count": len(elements_collection.tables), "images_count": len(elements_collection.inline_shapes)}
    for para in elements_collection.paragraphs:
        if para.text.strip(): # Only consider non-empty paragraphs
            p_info = {"text_snippet": para.text[:50].strip() + "..." if len(para.text.strip()) > 50 else para.text.strip()}
            p_info.update(get_paragraph_formatting_details(para))
            analysis["paragraphs_summary"].append(p_info)
    analysis["tables_details"] = []
    for table_idx, table in enumerate(elements_collection.tables):
        analysis["tables_details"].append({
            "index_in_hf": table_idx,
            "rows": len(table.rows),
            "cols": len(table.columns),
            "first_row_texts_snippet": [cell.text.strip()[:20] for cell in table.rows[0].cells][:3] if table.rows else []
        })
    return analysis

def analyze_document_structure(doc_path_or_file_like):
    """Analyzes a Word document and extracts its structure and formatting."""
    try:
        doc = Document(doc_path_or_file_like)
    except Exception as e:
        st.error(f"Error opening document: {e}")
        return {"error": f"Could not read document: {e}"}

    analysis = {
        "filename": os.path.basename(doc_path_or_file_like.name) if hasattr(doc_path_or_file_like, 'name') else "Unknown Filename",
        "sections_count": len(doc.sections),
        "page_setup_per_section": [],
        "headers_footers_per_section": [],
        "defined_styles_summary": {}, # Styles defined in the document
        "overall_element_counts": {
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "inline_images": len(doc.inline_shapes)
        },
        "toc_analysis": {"found_by_text": False, "heuristic_details": "Not detected by simple text search."},
        "document_tables_summary": [],
        "content_structure_preview": [] # First N paragraph styles/text
    }

    # Page Setup, Headers, Footers (per section)
    for i, section in enumerate(doc.sections):
        analysis["page_setup_per_section"].append({
            "section_index": i,
            "orientation": WD_ORIENT.to_xml(section.orientation).lower(),
            "page_width_cm": round(section.page_width.cm, 2) if section.page_width else None,
            "page_height_cm": round(section.page_height.cm, 2) if section.page_height else None,
            "top_margin_cm": round(section.top_margin.cm, 2) if section.top_margin else None,
            "bottom_margin_cm": round(section.bottom_margin.cm, 2) if section.bottom_margin else None,
            "left_margin_cm": round(section.left_margin.cm, 2) if section.left_margin else None,
            "right_margin_cm": round(section.right_margin.cm, 2) if section.right_margin else None,
        })
        header_analysis = analyze_header_footer_elements(section.header)
        footer_analysis = analyze_header_footer_elements(section.footer)
        analysis["headers_footers_per_section"].append({
            "section_index": i,
            "header": header_analysis,
            "footer": footer_analysis
        })

    # Styles Summary (paragraph styles defined in the document)
    # Prioritize common/important styles for detailed extraction
    priority_styles = ["Normal", "Body Text", "Title", "Subtitle"] + [f"Heading {i}" for i in range(1, 7)] + ["Caption", "List Paragraph"]
    
    for s in doc.styles:
        if s.type == WD_STYLE_TYPE.PARAGRAPH:
            try:
                style = doc.styles[s.name]
                font = style.font
                p_fmt = style.paragraph_format
                analysis["defined_styles_summary"][s.name] = {
                    "font_name": font.name,
                    "font_size_pt": round(font.size.pt,1) if font.size and font.size.pt is not None else None,
                    "font_bold": font.bold,
                    "font_italic": font.italic,
                    "font_underline": font.underline,
                    "alignment": WD_ALIGN_PARAGRAPH.to_xml(p_fmt.alignment).lower() if p_fmt.alignment is not None else 'left',
                    "space_before_pt": round(p_fmt.space_before.pt,1) if p_fmt.space_before and p_fmt.space_before.pt is not None else None,
                    "space_after_pt": round(p_fmt.space_after.pt,1) if p_fmt.space_after and p_fmt.space_after.pt is not None else None,
                    "line_spacing_val": round(p_fmt.line_spacing,2) if p_fmt.line_spacing is not None else None,
                    "built_in": style.builtin,
                }
            except KeyError: # Should not happen if iterating doc.styles
                analysis["defined_styles_summary"][s.name] = "Error fetching details for this style."
            except Exception: # Catch other potential issues with style properties
                 analysis["defined_styles_summary"][s.name] = "Could not reliably fetch all details for this style."


    # TOC Detection (simple heuristic: search for "Table of Contents" text or specific styles)
    for para_idx, para in enumerate(doc.paragraphs[:min(len(doc.paragraphs), 50)]): # Check first 50 paragraphs
        para_text_lower = para.text.lower().strip()
        if "table of contents" in para_text_lower or "contents" == para_text_lower :
            analysis["toc_analysis"]["found_by_text"] = True
            analysis["toc_analysis"]["heuristic_details"] = f"Found text suggestive of TOC in paragraph {para_idx+1}: '{para.text[:100].strip()}'"
            analysis["toc_analysis"]["style_of_toc_heading"] = para.style.name if para.style else "Default Paragraph Style"
            break
        if para.style and para.style.name and ("toc" in para.style.name.lower() or "TOC" in para.style.name):
            analysis["toc_analysis"]["found_by_text"] = True # Or a new key like found_by_style
            analysis["toc_analysis"]["heuristic_details"] = f"Found paragraph {para_idx+1} with a TOC-like style name: '{para.style.name}'"
            break


    # Document Tables Summary
    for i, table in enumerate(doc.tables):
        table_info = {
            "doc_table_index": i,
            "rows": len(table.rows),
            "cols": len(table.columns),
            "style_name": table.style.name if table.style else "Default Table Style",
            "autofit": table.autofit,
            "first_row_is_header_heuristic": False, # Basic check if first row is all bold
            "first_few_rows_text_snippet": []
        }
        if table.rows:
            first_row_cells = table.rows[0].cells
            is_header_bold = True
            if not first_row_cells: is_header_bold = False
            for cell_idx, cell in enumerate(first_row_cells):
                if not cell.paragraphs: # Empty cell
                    # If other cells in header are bold, an empty cell doesn't negate header
                    continue
                para = cell.paragraphs[0]
                # A cell is considered non-bold for header if it has runs and the first is not bold
                if para.runs and (para.runs[0].bold is None or para.runs[0].bold is False):
                    is_header_bold = False
                    break
                elif not para.runs: # Paragraph with no text/runs
                     # if other cells in header are bold, an empty cell doesn't negate header
                    pass
            table_info["first_row_is_header_heuristic"] = is_header_bold if first_row_cells else False

            for r_idx, row in enumerate(table.rows[:min(len(table.rows),3)]): # Snippet of first 3 rows
                row_texts = [cell.text.strip()[:30] for cell in row.cells[:min(len(row.cells),5)]] # First 5 cells, 30 chars
                table_info["first_few_rows_text_snippet"].append(row_texts)
        analysis["document_tables_summary"].append(table_info)

    # Content Structure Preview (first ~10 non-empty paragraphs)
    count = 0
    for para_idx, para in enumerate(doc.paragraphs):
        if count >= 10: break
        if para.text.strip():
            analysis["content_structure_preview"].append({
                "element_no_in_doc": para_idx,
                "type": "Paragraph",
                "style": para.style.name if para.style else "Default Paragraph Style",
                "text_snippet": para.text.strip()[:70] + ("..." if len(para.text.strip()) > 70 else "")
            })
            count+=1
            
    return analysis


# --- 2. Rule Engine (The template's analysis IS the rule set) ---

# --- 3. Document Validator ---
def compare_numerical_value(template_val, user_val, field_name, section_name="", tolerance=0.0):
    if template_val is None and user_val is None: return None
    if template_val is None or user_val is None: # One is None, other is not
         return f"[{section_name}] {field_name}: Template='{template_val}', Document='{user_val}' (one is not set)"
    if abs(template_val - user_val) > tolerance:
        return f"[{section_name}] {field_name}: Template='{template_val}', Document='{user_val}'"
    return None

def compare_text_value(template_val, user_val, field_name, section_name=""):
    if template_val is None and user_val is None: return None
    # Normalize common Nones vs empty strings or default values
    tv = template_val.lower().strip() if isinstance(template_val, str) else template_val
    uv = user_val.lower().strip() if isinstance(user_val, str) else user_val

    if tv == 'none' and uv is None: tv = None # Normalize 'None' string
    if uv == 'none' and tv is None: uv = None

    if tv != uv:
        return f"[{section_name}] {field_name}: Template='{template_val}', Document='{user_val}'"
    return None


def validate_document_against_template(user_doc_analysis, template_analysis, config):
    feedback = []
    priority_styles_to_check = config.get("priority_styles", ["Normal", "Heading 1", "Heading 2"])

    # 0. Error check
    if "error" in user_doc_analysis:
        feedback.append(f"Error processing user document: {user_doc_analysis['error']}")
        return feedback # Critical error, stop validation

    # 1. Section Count
    if template_analysis["sections_count"] != user_doc_analysis["sections_count"]:
         feedback.append(f"Section Count: Template has {template_analysis['sections_count']} section(s), Document has {user_doc_analysis['sections_count']}.")

    # 2. Page Setup (comparing first section for simplicity, or all if counts match)
    num_sections_to_compare = min(len(template_analysis["page_setup_per_section"]), len(user_doc_analysis["page_setup_per_section"]))
    for i in range(num_sections_to_compare if config.get("compare_all_sections", False) else min(1, num_sections_to_compare)): # Default to first section
        tpl_ps = template_analysis["page_setup_per_section"][i]
        usr_ps = user_doc_analysis["page_setup_per_section"][i]
        section_label = f"Section {i} Page Setup"

        diff = compare_text_value(tpl_ps.get("orientation"), usr_ps.get("orientation"), "Orientation", section_label)
        if diff: feedback.append(diff)
        
        margin_tolerance = config.get("margin_tolerance_cm", 0.1)
        for dim in ["page_width_cm", "page_height_cm", "top_margin_cm", "bottom_margin_cm", "left_margin_cm", "right_margin_cm"]:
            diff = compare_numerical_value(tpl_ps.get(dim), usr_ps.get(dim), dim.replace("_", " ").title(), section_label, tolerance=margin_tolerance)
            if diff: feedback.append(diff)

    # 3. Headers & Footers (basic comparison for presence and element counts for first section or all)
    num_hf_sections_to_compare = min(len(template_analysis["headers_footers_per_section"]), len(user_doc_analysis["headers_footers_per_section"]))
    for i in range(num_hf_sections_to_compare if config.get("compare_all_sections", False) else min(1, num_hf_sections_to_compare)):
        tpl_hf_sec = template_analysis["headers_footers_per_section"][i]
        usr_hf_sec = user_doc_analysis["headers_footers_per_section"][i]

        for hf_type in ["header", "footer"]:
            tpl_hf_content = tpl_hf_sec[hf_type]
            usr_hf_content = usr_hf_sec[hf_type]
            hf_label = f"Section {i} {hf_type.title()}"

            tpl_has_content = bool(tpl_hf_content["paragraphs_summary"] or tpl_hf_content["tables_count"] > 0)
            usr_has_content = bool(usr_hf_content["paragraphs_summary"] or usr_hf_content["tables_count"] > 0)

            if tpl_has_content and not usr_has_content:
                feedback.append(f"{hf_label}: Template has content, but Document's appears empty or significantly different.")
            elif not tpl_has_content and usr_has_content:
                feedback.append(f"{hf_label}: Template has no significant content, but Document has content.")
            
            if tpl_hf_content["tables_count"] != usr_hf_content["tables_count"]:
                 feedback.append(f"{hf_label} Tables: Template has {tpl_hf_content['tables_count']}, Document has {usr_hf_content['tables_count']}.")
            # Further detail: compare paragraph styles/text snippets if desired (more complex)


    # 4. Styles Summary (focus on priority styles or all defined in template)
    styles_to_check_in_template = template_analysis["defined_styles_summary"]
    
    for style_name, tpl_style_details in styles_to_check_in_template.items():
        if not config.get("check_all_styles", False) and style_name not in priority_styles_to_check:
            if not (tpl_style_details.get("built_in", False) == False and config.get("check_custom_styles", True)): # If not priority, only check custom styles if flag is true
                 continue # Skip non-priority styles if not checking all or not a custom style to check

        usr_style_details = user_doc_analysis["defined_styles_summary"].get(style_name)
        
        if isinstance(tpl_style_details, str): # Error fetching template style
            feedback.append(f"Info: Could not analyze style '{style_name}' in template: {tpl_style_details}")
            continue

        if not usr_style_details:
            if style_name in priority_styles_to_check or not tpl_style_details.get("built_in", True): # Flag missing priority or custom styles
                feedback.append(f"Style Missing: Style '{style_name}' is defined in template but not found in user document.")
            continue
        if isinstance(usr_style_details, str): # Error fetching user style
            feedback.append(f"Info: Could not analyze style '{style_name}' in user document: {usr_style_details}")
            continue

        # Compare individual properties of the style
        style_prop_tolerance = config.get("style_font_size_tolerance_pt", 0.5)
        style_spacing_tolerance = config.get("style_spacing_tolerance_pt", 1.0)

        props_map = { # field_name: (comparison_func, tolerance_key_in_config_if_numerical)
            "font_name": (compare_text_value, None),
            "font_size_pt": (compare_numerical_value, style_prop_tolerance),
            "font_bold": (compare_text_value, None), # Boolean will be compared as text 'True'/'False'
            "font_italic": (compare_text_value, None),
            "alignment": (compare_text_value, None),
            "space_before_pt": (compare_numerical_value, style_spacing_tolerance),
            "space_after_pt": (compare_numerical_value, style_spacing_tolerance),
            "line_spacing_val": (compare_numerical_value, 0.05),
        }
        
        for prop, (comp_func, tolerance) in props_map.items():
            tpl_val = tpl_style_details.get(prop)
            usr_val = usr_style_details.get(prop)
            
            # Handle boolean conversion for comparison if needed for bool props
            if prop in ["font_bold", "font_italic"]:
                tpl_val = str(tpl_val) if tpl_val is not None else None
                usr_val = str(usr_val) if usr_val is not None else None

            args = [tpl_val, usr_val, prop.replace("_", " ").title(), f"Style '{style_name}'"]
            if comp_func == compare_numerical_value:
                args.append(tolerance)
            
            diff = comp_func(*args)
            if diff: feedback.append(diff)

    # Check for styles in user doc not in template's defined styles (potential 'extra' styles if they are custom)
    if config.get("report_extra_styles", True):
        for style_name, usr_style_details in user_doc_analysis["defined_styles_summary"].items():
            if style_name not in styles_to_check_in_template and not usr_style_details.get("built_in", True) and isinstance(usr_style_details, dict):
                feedback.append(f"Additional Custom Style: Document defines custom style '{style_name}' which is not in the template's defined styles list.")

    # 5. TOC
    if template_analysis["toc_analysis"]["found_by_text"] and not user_doc_analysis["toc_analysis"]["found_by_text"]:
        feedback.append("Table of Contents: Expected (by text heuristic), but not detected in user document.")
    elif not template_analysis["toc_analysis"]["found_by_text"] and user_doc_analysis["toc_analysis"]["found_by_text"]:
        feedback.append("Table of Contents: Not detected in template (by text heuristic), but present in user document.")

    # 6. Overall Element Counts
    tpl_counts = template_analysis["overall_element_counts"]
    usr_counts = user_doc_analysis["overall_element_counts"]
    
    count_tolerance_percent = config.get("element_count_tolerance_percent", 20) # 20% tolerance

    for elem_type in ["tables", "inline_images"]:
        tpl_c = tpl_counts.get(elem_type, 0)
        usr_c = usr_counts.get(elem_type, 0)
        allowed_diff = (tpl_c * count_tolerance_percent / 100.0)
        if abs(tpl_c - usr_c) > max(2, allowed_diff) : # Allow at least 2 difference or percentage
            feedback.append(f"{elem_type.replace('_', ' ').title()} Count: Template has ~{tpl_c}, Document has {usr_c}.")

    # 7. Table Structure (comparing first table more closely if present)
    if template_analysis["document_tables_summary"] and user_doc_analysis["document_tables_summary"]:
        tpl_table0 = template_analysis["document_tables_summary"][0]
        usr_table0 = user_doc_analysis["document_tables_summary"][0]
        table_label = "First Document Table"

        if tpl_table0["cols"] != usr_table0["cols"]:
            feedback.append(f"{table_label} Column Count: Template has {tpl_table0['cols']}, Document has {usr_table0['cols']}.")
        if tpl_table0["first_row_is_header_heuristic"] != usr_table0["first_row_is_header_heuristic"]:
            feedback.append(f"{table_label} Header Row (Bold Heuristic): Template suggests '{tpl_table0['first_row_is_header_heuristic']}', Document suggests '{usr_table0['first_row_is_header_heuristic']}'.")
    elif template_analysis["document_tables_summary"] and not user_doc_analysis["document_tables_summary"]:
        feedback.append("Tables: Template has table details, but none found or analyzed in user document's main body.")

    if not feedback:
        feedback.append("No major formatting discrepancies detected based on the current rule set and configuration. (Note: This is not an exhaustive check of all Word features).")

    return feedback

# --- 4. Streamlit UI ---
st.set_page_config(layout="wide", page_title="Word Doc Formatter")
st.title("📄 Word Document Format Analyzer & Validator")
st.markdown("""
This tool analyzes a template Word document to establish a formatting baseline (rules).
It then validates other Word documents against these rules, providing feedback on deviations.
""")

# --- Configuration Sidebar ---
st.sidebar.header("⚙️ Validation Configuration")
# Add more config options as needed
config = {
    "priority_styles": st.sidebar.text_input("Priority Styles to Check (comma-separated)", "Normal,Heading 1,Heading 2,Heading 3,Body Text,Caption,List Paragraph").split(','),
    "margin_tolerance_cm": st.sidebar.slider("Margin Tolerance (cm)", 0.0, 0.5, 0.1, 0.01),
    "style_font_size_tolerance_pt": st.sidebar.slider("Style Font Size Tolerance (pt)", 0.0, 2.0, 0.5, 0.1),
    "style_spacing_tolerance_pt": st.sidebar.slider("Style Spacing Tolerance (pt)", 0.0, 5.0, 1.0, 0.5),
    "element_count_tolerance_percent": st.sidebar.slider("Element Count Tolerance (%)", 0, 100, 20, 5),
    "check_all_styles": st.sidebar.checkbox("Check All Defined Styles (not just priority/custom)", False),
    "check_custom_styles": st.sidebar.checkbox("Check Custom Styles (if not in priority list)", True),
    "report_extra_styles": st.sidebar.checkbox("Report Extra Custom Styles in User Doc", True),
    "compare_all_sections": st.sidebar.checkbox("Compare All Sections (Page Setup, Headers/Footers)", False, help="If unchecked, only the first section is compared for these elements."),
}


# --- Template Section ---
st.sidebar.header("📐 Template Document")
template_file = st.sidebar.file_uploader("1. Upload Template (.docx)", type=["docx"], key="template_uploader")

if 'template_analysis' not in st.session_state:
    st.session_state.template_analysis = None
if 'template_name' not in st.session_state:
    st.session_state.template_name = None

if template_file:
    if st.sidebar.button("Analyze Template Document", key="analyze_template_btn", use_container_width=True):
        st.session_state.template_name = template_file.name
        with st.spinner(f"Analyzing template '{template_file.name}'..."):
            bytes_data = template_file.getvalue()
            file_like_object = io.BytesIO(bytes_data)
            file_like_object.name = template_file.name # Critical for filename in analysis
            st.session_state.template_analysis = analyze_document_structure(file_like_object)

            if "error" in st.session_state.template_analysis:
                st.error(f"Error analyzing template: {st.session_state.template_analysis['error']}")
                st.session_state.template_analysis = None # Clear on error
            else:
                st.success(f"Template '{template_file.name}' analyzed successfully! This is now the rule set.")

st.divider()
st.header("📋 Template Analysis (Rule Set)")
if st.session_state.template_analysis:
    st.info(f"**Current Template:** `{st.session_state.template_name}`. All user documents will be validated against this.")
    analysis_data = st.session_state.template_analysis
    if "error" not in analysis_data:
        with st.expander("View Full Template Analysis (JSON Data)", expanded=False):
            st.json(analysis_data, expanded=False)

        # Display a summary of the template
        col1, col2 = st.columns(2)
        with col1:
            st.subheader("📄 Page & Structure:")
            if analysis_data["page_setup_per_section"]:
                ps0 = analysis_data["page_setup_per_section"][0]
                st.markdown(f"- **Orientation (Sec 0):** `{ps0.get('orientation', 'N/A')}`")
                st.markdown(f"- **Margins (T,B,L,R cm - Sec 0):** `{ps0.get('top_margin_cm','N/A')}, {ps0.get('bottom_margin_cm','N/A')}, {ps0.get('left_margin_cm','N/A')}, {ps0.get('right_margin_cm','N/A')}`")
            st.markdown(f"- **Total Sections:** `{analysis_data.get('sections_count', 'N/A')}`")
            if analysis_data["headers_footers_per_section"]:
                hf0 = analysis_data["headers_footers_per_section"][0]
                st.markdown(f"- **Header (Sec 0):** {'Content detected' if hf0['header']['paragraphs_summary'] or hf0['header']['tables_count'] > 0 else 'No significant content detected'}")
                st.markdown(f"- **Footer (Sec 0):** {'Content detected' if hf0['footer']['paragraphs_summary'] or hf0['footer']['tables_count'] > 0 else 'No significant content detected'}")
            st.markdown(f"- **Table of Contents:** {'Detected (heuristic)' if analysis_data['toc_analysis']['found_by_text'] else 'Not detected (heuristic)'}")
            st.markdown(f"- **Total Tables (doc body):** `{analysis_data['overall_element_counts']['tables']}`")
            st.markdown(f"- **Total Inline Images (doc body):** `{analysis_data['overall_element_counts']['inline_images']}`")

        with col2:
            st.subheader("💅 Key Monitored Styles:")
            monitored_styles_in_template = 0
            for style_name in config["priority_styles"]:
                details = analysis_data["defined_styles_summary"].get(style_name)
                if isinstance(details, dict):
                    monitored_styles_in_template +=1
                    font_str = f"{details.get('font_name', 'N/A')}, {details.get('font_size_pt', 'N/A')}pt"
                    font_str += " B" if details.get('font_bold') else ""
                    font_str += " I" if details.get('font_italic') else ""
                    st.markdown(f"- **`{style_name}`**: {font_str} ({details.get('alignment', 'N/A')})")
                elif style_name in analysis_data["defined_styles_summary"] : # Style was checked but maybe had an issue or not found
                    st.markdown(f"- **`{style_name}`**: {details if isinstance(details, str) else 'Not found in template styles.'}")

            if monitored_styles_in_template == 0 and config["priority_styles"]:
                st.caption("None of the specified priority styles were found with details in the template.")
            elif not config["priority_styles"]:
                 st.caption("No priority styles specified in configuration.")

    else: # Error in analysis_data
        st.warning(f"Could not display template analysis: {analysis_data.get('error', 'Unknown error')}")
else:
    st.info("⬆️ Upload a template Word document and click 'Analyze Template Document' (in the sidebar) to establish the formatting rules.")


# --- User Documents Validation Section ---
st.divider()
st.header("🔎 Validate User Documents")

if st.session_state.template_analysis and "error" not in st.session_state.template_analysis:
    user_doc_files = st.file_uploader("2. Upload User Documents to Validate (.docx)", type=["docx"], accept_multiple_files=True, key="user_docs_uploader")

    if user_doc_files:
        for user_doc_file in user_doc_files:
            st.subheader(f"Validation Feedback for: `{user_doc_file.name}`")
            with st.spinner(f"Analyzing and validating '{user_doc_file.name}'..."):
                user_bytes_data = user_doc_file.getvalue()
                user_file_like_object = io.BytesIO(user_bytes_data)
                user_file_like_object.name = user_doc_file.name # Add .name attribute

                user_doc_analysis = analyze_document_structure(user_file_like_object)

                if "error" in user_doc_analysis:
                    st.error(f"Error analyzing document '{user_doc_file.name}': {user_doc_analysis['error']}")
                    continue # Skip to next file

                feedback_messages = validate_document_against_template(user_doc_analysis, st.session_state.template_analysis, config)

                if feedback_messages:
                    for msg_idx, msg in enumerate(feedback_messages):
                        is_major = "Mismatch" in msg or "Missing" in msg or "Error" in msg or "Discrepancy" in msg or "Count:" in msg or "Additional" in msg or "different" in msg
                        is_info = "Info:" in msg or "suggests" in msg
                        
                        if msg.startswith("No major formatting discrepancies"):
                            st.success(f"✅ {msg}", icon="🎉")
                        elif is_major:
                            st.warning(f"⚠️ {msg}")
                        elif is_info:
                            st.info(f"ℹ️ {msg}")
                        else: 
                            st.info(f"➡️ {msg}") # Default for other messages
                else: # Should not happen due to the default "No major discrepancies" message
                    st.info("No feedback messages generated. This might indicate an issue if discrepancies were expected.")

                with st.expander("View Full Analysis of this User Document (JSON)"):
                    st.json(user_doc_analysis)
            st.markdown("---") # Separator between validated documents
    elif st.session_state.template_analysis: # Template loaded, but no user docs yet
        st.info("⬆️ Upload one or more user documents to validate them against the loaded template.")
else: # No template loaded/analyzed yet
    st.warning("⚠️ Please upload and analyze a template document first before attempting to validate user documents.")

st.sidebar.divider()
st.sidebar.markdown("Built by an AI Assistant.")
st.sidebar.caption("This tool offers basic structural and style validation. It may not capture all nuances of Word document formatting.")
