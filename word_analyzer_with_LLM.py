import streamlit as st
from docx import Document
from docx.shared import Pt, Cm 
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
import io
import os
import json # For LLM response parsing

# Attempt to import LLM an related libraries
try:
    import google.generativeai as genai
    LLM_AVAILABLE = True
except ImportError:
    LLM_AVAILABLE = False
    st.warning("The 'google-generativeai' package is not installed. LLM features will be disabled. Install with: pip install google-generativeai")

# --- 1. Document Analyzer (Structural and Stylistic) ---
def get_paragraph_formatting_details(para):
    """Extracts formatting details from a paragraph object."""
    details = {}
    if para.style and para.style.name:
        details['style_name'] = para.style.name

    if para.runs:
        run = para.runs[0] # Dominant formatting from the first run
        font = run.font
        details['font_name'] = font.name
        details['font_size_pt'] = round(font.size.pt, 1) if font.size and font.size.pt is not None else None
        details['font_bold'] = font.bold
        details['font_italic'] = font.italic
        details['font_underline'] = font.underline
        details['font_color_rgb'] = str(font.color.rgb) if font.color and font.color.rgb else None

    p_format = para.paragraph_format
    details['alignment'] = WD_ALIGN_PARAGRAPH.to_xml(p_format.alignment).lower() if p_format.alignment is not None else 'left'
    details['space_before_pt'] = round(p_format.space_before.pt, 1) if p_format.space_before and p_format.space_before.pt is not None else None
    details['space_after_pt'] = round(p_format.space_after.pt, 1) if p_format.space_after and p_format.space_after.pt is not None else None
    details['line_spacing_val'] = round(p_format.line_spacing, 2) if p_format.line_spacing is not None else None
    details['left_indent_cm'] = round(p_format.left_indent.cm, 2) if p_format.left_indent and p_format.left_indent.cm is not None else None
    details['first_line_indent_cm'] = round(p_format.first_line_indent.cm, 2) if p_format.first_line_indent and p_format.first_line_indent.cm is not None else None
    return details

def analyze_header_footer_elements(elements_collection):
    """Analyzes paragraphs and tables within a header/footer collection."""
    analysis = {"paragraphs_summary": [], "tables_count": len(elements_collection.tables), "images_count": len(elements_collection.inline_shapes)}
    for para in elements_collection.paragraphs:
        if para.text.strip():
            p_info = {"text_snippet": para.text[:50].strip() + "..." if len(para.text.strip()) > 50 else para.text.strip()}
            p_info.update(get_paragraph_formatting_details(para))
            analysis["paragraphs_summary"].append(p_info)
    # Add more details for tables in headers/footers if needed
    return analysis

def analyze_document_structure(doc_path_or_file_like, is_template=False):
    """
    Analyzes a Word document for structure, styles, and optionally, section descriptions (if is_template=True).
    """
    try:
        doc = Document(doc_path_or_file_like)
    except Exception as e:
        return {"error": f"Could not read document: {e}"}

    analysis = {
        "filename": os.path.basename(doc_path_or_file_like.name) if hasattr(doc_path_or_file_like, 'name') else "UploadedFile",
        "sections_count": len(doc.sections),
        "page_setup_per_section": [],
        "headers_footers_per_section": [],
        "defined_styles_summary": {},
        "overall_element_counts": {
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "inline_images": len(doc.inline_shapes)
        },
        "toc_analysis": {"found_by_text": False, "heuristic_details": "Not detected by simple text search."},
        "document_tables_summary": [],
        "content_structure_preview": [],
        "sections_with_descriptions": [] # For template: to store {title, description_from_template}
    }

    # Page Setup, Headers, Footers (per section)
    for i, section in enumerate(doc.sections):
        analysis["page_setup_per_section"].append({
            "section_index": i,
            "orientation": WD_ORIENT.to_xml(section.orientation).lower(),
            "page_width_cm": round(section.page_width.cm, 2) if section.page_width else None,
            "page_height_cm": round(section.page_height.cm, 2) if section.page_height else None,
            # ... (add other margin details as in previous versions) ...
        })
        header_analysis = analyze_header_footer_elements(section.header)
        footer_analysis = analyze_header_footer_elements(section.footer)
        analysis["headers_footers_per_section"].append({
            "section_index": i, "header": header_analysis, "footer": footer_analysis
        })

    # Styles Summary
    for s in doc.styles:
        if s.type == WD_STYLE_TYPE.PARAGRAPH:
            try:
                # ... (style detail extraction as in previous versions) ...
                analysis["defined_styles_summary"][s.name] = {
                    "font_name": s.font.name, "font_size_pt": s.font.size.pt if s.font.size else None, # Simplified
                    "built_in": s.builtin
                }
            except Exception:
                 analysis["defined_styles_summary"][s.name] = "Error fetching details."
    
    # TOC Detection (heuristic)
    # ... (TOC detection logic as in previous versions) ...

    # Document Tables Summary
    # ... (Table analysis logic as in previous versions) ...

    # If it's a template, try to extract section titles and their descriptions
    if is_template:
        # Heuristic: "Heading 1" style marks a section title.
        # The paragraph(s) immediately following, not being another major heading, form the description.
        # This needs to be robust or configurable based on the template's specific conventions.
        current_section_title = None
        current_description_paras_texts = []
        paras_list = list(doc.paragraphs) # Make it a list for easier lookahead

        for i, para in enumerate(paras_list):
            para_text_stripped = para.text.strip()
            style_name = para.style.name if para.style else ""

            is_section_heading_style = style_name == st.session_state.get("template_section_heading_style", "Heading 1")

            if is_section_heading_style and para_text_stripped:
                # If there was a previous section being tracked, save it
                if current_section_title and current_description_paras_texts:
                    analysis["sections_with_descriptions"].append({
                        "title": current_section_title,
                        "description_from_template": "\n".join(current_description_paras_texts)
                    })
                current_section_title = para_text_stripped
                current_description_paras_texts = [] # Reset for the new section
            elif current_section_title: # We are inside a section, collecting its description
                # Stop collecting if we hit another major heading or a style that indicates end of description
                # This is a simplified condition; more complex logic might be needed.
                next_para_is_new_section_heading = False
                if (i > 0 and style_name == st.session_state.get("template_section_heading_style", "Heading 1")):
                     next_para_is_new_section_heading = True


                if not next_para_is_new_section_heading and para_text_stripped:
                    # Add to current description if it's not another heading of same/higher level
                    # and not a style that explicitly breaks a description block.
                    # For simplicity, we add non-empty paragraphs until the next designated section heading.
                    current_description_paras_texts.append(para_text_stripped)
                elif next_para_is_new_section_heading : # hit a new section heading, so the previous description ends
                    if current_section_title and current_description_paras_texts:
                         analysis["sections_with_descriptions"].append({
                            "title": current_section_title,
                            "description_from_template": "\n".join(current_description_paras_texts)
                        })
                    current_section_title = para_text_stripped # Start new section
                    current_description_paras_texts = []


        # Add the last collected section description, if any
        if current_section_title and current_description_paras_texts:
            analysis["sections_with_descriptions"].append({
                "title": current_section_title,
                "description_from_template": "\n".join(current_description_paras_texts)
            })
    return analysis

# --- 2. User Document Section Content Extractor ---
def extract_user_section_content(user_doc_obj, section_title_from_template, config):
    """
    Extracts text content from a user document for a given section title.
    Stops at the next heading of the same or higher level as specified in config.
    """
    content_paragraphs = []
    in_target_section = False
    # Section heading styles that would break the current section's content gathering.
    # Typically, the template's main section heading style, and potentially higher ones like "Heading 1" if current is "Heading 2".
    # This needs to be more sophisticated if there's a deep hierarchy.
    breaking_heading_styles = [config.get("template_section_heading_style", "Heading 1")]
    # If template_section_heading_style is e.g. "Heading 2", then "Heading 1" should also break.
    try:
        main_level = int(config.get("template_section_heading_style", "Heading 1").split(" ")[-1])
        for i in range(1, main_level):
            breaking_heading_styles.append(f"Heading {i}")
    except: # Not a numbered heading style
        pass
    
    # Normalize to ensure uniqueness and handle potential user input variations
    breaking_heading_styles = list(set(s.strip() for s in breaking_heading_styles if s and s.strip()))


    for para_idx, para in enumerate(user_doc_obj.paragraphs):
        para_text_stripped = para.text.strip()
        current_style_name = para.style.name.strip() if para.style and para.style.name else ""

        if not in_target_section:
            # Case-insensitive matching for section titles is often more robust
            if para_text_stripped.lower() == section_title_from_template.lower().strip():
                in_target_section = True
                # Optionally, skip adding the title itself to the content:
                # continue 
        elif in_target_section:
            # Check if this paragraph marks the beginning of a new, breaking section
            if current_style_name in breaking_heading_styles and para_text_stripped.lower() != section_title_from_template.lower().strip():
                break # End of current section's content

            if para_text_stripped:
                content_paragraphs.append(para_text_stripped)
    
    return "\n".join(content_paragraphs)


# --- 3. LLM API Integration ---
@st.cache_data(show_spinner=False) # Cache LLM calls for same inputs
def get_llm_content_relevance(_llm_model_instance, template_description, user_section_content, llm_model_name_for_cache_key):
    """
    Calls the LLM to assess content relevance.
    _llm_model_instance is the configured genai.GenerativeModel object.
    llm_model_name_for_cache_key is just to make Streamlit's caching work if model changes.
    """
    if not user_section_content.strip():
        return {"relevance_score": 0, "explanation": "Section appears empty or was not found in the user document.", "error": None}
    if not template_description.strip():
        return {"relevance_score": None, "explanation": "Template description for this section is empty. Cannot assess relevance.", "error": None}

    # Simple truncation for token safety. Consider more advanced chunking for production.
    max_chars_user_content = 7000  # Gemini models have larger context windows
    max_chars_description = 2000

    prompt = f"""
    You are an expert document content analyzer. Your task is to evaluate how well a user-provided text section aligns with a template description that outlines what the section should ideally contain.

    Template Section Description:
    ---
    {template_description[:max_chars_description]}
    ---

    User Document Section Content:
    ---
    {user_section_content[:max_chars_user_content]}
    ---

    Please perform the following:
    1.  Carefully understand the "Template Section Description" to identify the required topics, information, and the overall purpose of this section.
    2.  Thoroughly analyze the "User Document Section Content" to determine how effectively it addresses each requirement from the description.
    3.  Provide a concise "explanation" of your findings. Specifically mention:
        - What aspects of the description are well-covered in the user's content.
        - What aspects are partially covered or addressed vaguely.
        - What required aspects are missing entirely.
        - Any content in the user's section that seems irrelevant to the template description.
    4.  Assign a "relevance_score" as an integer between 0 (not relevant at all / key information missing) and 10 (perfectly relevant and comprehensive).

    IMPORTANT: Return your response ONLY as a valid JSON object with exactly two keys: "explanation" (string) and "relevance_score" (integer).
    Do not include any other text, introductory phrases, or markdown formatting outside of this JSON object.

    Example of the required JSON output format:
    {{
        "explanation": "The user's content comprehensively covers the project goals and methodology as described. However, it omits the 'expected outcomes' section and includes a lengthy discussion on unrelated prior work which was not requested.",
        "relevance_score": 6
    }}
    """
    try:
        response = _llm_model_instance.generate_content(prompt)
        
        # Clean and parse the LLM response
        raw_text_response = response.text.strip()
        # Remove potential markdown code block delimiters
        if raw_text_response.startswith("```json"):
            raw_text_response = raw_text_response[7:-3].strip()
        elif raw_text_response.startswith("```"):
             raw_text_response = raw_text_response[3:-3].strip()

        parsed_llm_result = json.loads(raw_text_response)
        
        # Validate expected keys
        if "explanation" not in parsed_llm_result or "relevance_score" not in parsed_llm_result:
            raise ValueError("LLM JSON response missing required keys ('explanation', 'relevance_score').")
        if not isinstance(parsed_llm_result["relevance_score"], int):
             # Try to convert if it's a stringified int, else raise error
             try:
                 parsed_llm_result["relevance_score"] = int(parsed_llm_result["relevance_score"])
             except ValueError:
                 raise ValueError("'relevance_score' must be an integer.")


        parsed_llm_result["error"] = None # Add error key for consistency
        return parsed_llm_result

    except json.JSONDecodeError as je:
        error_msg = f"LLM response was not valid JSON: {str(je)}. Raw response: '{raw_text_response[:200]}...'"
        return {"relevance_score": None, "explanation": error_msg, "error": error_msg}
    except Exception as e:
        error_msg = f"LLM API call or response processing failed: {str(e)}"
        return {"relevance_score": None, "explanation": error_msg, "error": error_msg}

# --- 4. Document Validator ---
def validate_document_against_template(user_doc_analysis, template_analysis, config, llm_model_instance_for_validation=None):
    feedback = []
    # --- Structural and Stylistic Checks (Simplified for brevity, expand as needed) ---
    if template_analysis["sections_count"] != user_doc_analysis["sections_count"]:
         feedback.append(f"⚠️ Section Count: Template has {template_analysis['sections_count']}, Document has {user_doc_analysis['sections_count']}.")
    # ... Add more stylistic comparisons for fonts, margins, etc., from previous versions ...

    # --- LLM Content Relevance Check ---
    if config.get("enable_llm_validation") and template_analysis.get("sections_with_descriptions"):
        feedback.append("--- 🧠 LLM Content Relevance Analysis ---")
        
        if not llm_model_instance_for_validation: # Check if LLM model was successfully initialized
            feedback.append("⚠️ LLM Model not available (check API key & setup). Skipping content relevance.")
        else:
            # To extract content, we need the user's Document object.
            # It's assumed "raw_bytes" of the user document were stored in user_doc_analysis.
            if "raw_bytes" not in user_doc_analysis:
                feedback.append("⚠️ 'raw_bytes' of user document not found. Cannot perform LLM section analysis.")
            else:
                user_doc_obj_for_llm = Document(io.BytesIO(user_doc_analysis["raw_bytes"]))

                for section_detail in template_analysis["sections_with_descriptions"]:
                    template_title = section_detail["title"]
                    template_desc_text = section_detail["description_from_template"]
                    
                    feedback_prefix = f"Section '{template_title}': "
                    with st.spinner(f"LLM evaluating content for section: '{template_title}'..."):
                        user_content_text = extract_user_section_content(user_doc_obj_for_llm, template_title, config)
                        
                        # Pass the actual model name to the cache key for get_llm_content_relevance
                        llm_result = get_llm_content_relevance(
                            llm_model_instance_for_validation, 
                            template_desc_text, 
                            user_content_text,
                            config.get("llm_model_name", "default_model_key") # For caching
                        )
                    
                    score = llm_result.get('relevance_score')
                    explanation = llm_result.get('explanation', 'No explanation provided.')
                    error = llm_result.get('error')

                    if error:
                        feedback.append(f"❌ {feedback_prefix}LLM Error - {explanation}")
                    elif score is not None:
                        emoji = "✅" if score >= 7 else ("🔶" if score >= 4 else "❌")
                        feedback.append(f"{emoji} {feedback_prefix}Relevance Score: {score}/10")
                        feedback.append(f"   └── LLM Explanation: {explanation}")
                    else: # Score is None but no specific error message from LLM function itself
                        feedback.append(f"❓ {feedback_prefix}Could not determine relevance. LLM Explanation: {explanation}")
    elif config.get("enable_llm_validation"):
        feedback.append("ℹ️ LLM validation enabled, but no sections with descriptions were found in the template to analyze.")


    if not feedback:
        feedback.append("✅ No major discrepancies detected based on current rules (structural checks might be minimal in this version).")
    return feedback

# --- 5. Streamlit UI ---
st.set_page_config(layout="wide", page_title="Word Document Analyzer & Validator")
st.title("📄 Word Document Format & Content Validator")
st.markdown("""
Upload a template Word document to define formatting rules and section content expectations.
Then, upload user documents to validate them against the template.
LLM-powered content relevance checking can be enabled in the sidebar.
""")

# --- Initialize Session State ---
if 'template_analysis' not in st.session_state:
    st.session_state.template_analysis = None
if 'template_name' not in st.session_state:
    st.session_state.template_name = None
if 'llm_api_key' not in st.session_state: # Store API key in session state
    st.session_state.llm_api_key = ""
if 'llm_model_name' not in st.session_state:
    st.session_state.llm_model_name = "gemini-1.5-flash-latest" # Default model
if 'enable_llm_validation_cb' not in st.session_state:
    st.session_state.enable_llm_validation_cb = False
if 'template_section_heading_style' not in st.session_state:
    st.session_state.template_section_heading_style = "Heading 1"

# --- Global LLM Model Instance (initialized if API key is valid) ---
# This allows initializing once per session if key is stable.
llm_model_global_instance = None
if LLM_AVAILABLE and st.session_state.get("enable_llm_validation_cb") and st.session_state.get("llm_api_key"):
    try:
        genai.configure(api_key=st.session_state.llm_api_key)
        llm_model_global_instance = genai.GenerativeModel(st.session_state.llm_model_name)
        # Simple test to see if model is accessible, can be removed or made more robust
        # llm_model_global_instance.generate_content("test", generation_config=genai.types.GenerationConfig(max_output_tokens=5))
        if 'llm_init_success' not in st.session_state or not st.session_state.llm_init_success:
            st.sidebar.success(f"LLM '{st.session_state.llm_model_name}' initialized.")
            st.session_state.llm_init_success = True
    except Exception as e:
        st.sidebar.error(f"LLM Init Error: {str(e)[:100]}...")
        llm_model_global_instance = None
        st.session_state.llm_init_success = False
elif st.session_state.get("enable_llm_validation_cb") and not st.session_state.get("llm_api_key"):
     st.sidebar.warning("LLM validation enabled, but API key is missing.")
     st.session_state.llm_init_success = False


# --- Sidebar ---
with st.sidebar:
    st.header("📐 Template Setup")
    template_file = st.file_uploader("1. Upload Template (.docx)", type=["docx"], key="template_uploader")
    st.session_state.template_section_heading_style = st.text_input(
        "Style for Section Headings in Template", 
        value=st.session_state.template_section_heading_style,
        help="The exact style name (e.g., 'Heading 1', 'My Section Style') used in your template to denote main section titles whose descriptions should be extracted."
    )

    if template_file:
        if st.button("Analyze Template Document", key="analyze_template_btn", use_container_width=True):
            st.session_state.template_name = template_file.name
            with st.spinner(f"Analyzing template '{template_file.name}'..."):
                bytes_data = template_file.getvalue()
                file_like_object = io.BytesIO(bytes_data)
                file_like_object.name = template_file.name
                # Pass is_template=True to enable section/description extraction
                st.session_state.template_analysis = analyze_document_structure(file_like_object, is_template=True)

                if "error" in st.session_state.template_analysis:
                    st.error(f"Error analyzing template: {st.session_state.template_analysis['error']}")
                    st.session_state.template_analysis = None
                else:
                    st.success(f"Template '{template_file.name}' analyzed.")
                    if not st.session_state.template_analysis.get("sections_with_descriptions"):
                        st.warning(f"No section descriptions extracted from template. Ensure section titles use the style '{st.session_state.template_section_heading_style}' and are followed by their description paragraphs.")
    
    st.divider()
    st.header("⚙️ Validation Configuration")
    config = { # Store config in a dictionary
        "priority_styles": st.text_input("Priority Styles (comma-sep)", "Normal,Heading 1,Heading 2").split(','),
        # Add other structural/stylistic config options here if needed
    }

    st.divider()
    st.header("🧠 LLM Content Validation")
    if not LLM_AVAILABLE:
        st.error("'google-generativeai' not installed. LLM features disabled.")
    else:
        config["enable_llm_validation"] = st.checkbox(
            "Enable LLM Content Relevance Check", 
            value=st.session_state.enable_llm_validation_cb, 
            key="enable_llm_cb_main_controller" # Use a new key to avoid conflict if it was set by direct assignment before
        )
        st.session_state.enable_llm_validation_cb = config["enable_llm_validation"] # Sync back

        if config["enable_llm_validation"]:
            st.session_state.llm_api_key = st.text_input(
                "LLM API Key (Google Gemini)", type="password", 
                value=st.session_state.llm_api_key,
                help="Your API key from Google AI Studio or Vertex AI."
            )
            # Update LLM Model Name in session state from selectbox
            st.session_state.llm_model_name = st.selectbox(
               "LLM Model",
               ["gemini-1.5-flash-latest", "gemini-pro", "gemini-1.0-pro"], # Common models
               index=["gemini-1.5-flash-latest", "gemini-pro", "gemini-1.0-pro"].index(st.session_state.llm_model_name) if st.session_state.llm_model_name in ["gemini-1.5-flash-latest", "gemini-pro", "gemini-1.0-pro"] else 0,
               key="llm_model_select_sidebar"
            )
            config["llm_model_name"] = st.session_state.llm_model_name # Pass to config for validation function
            config["template_section_heading_style"] = st.session_state.template_section_heading_style # Pass to config for section extraction

            if st.button("Clear LLM Cache", help="Clears cached LLM responses. Useful if you've updated prompts or expect different results for same input."):
                get_llm_content_relevance.clear()
                st.success("LLM cache cleared.")


# --- Main Area for Display and Validation ---
st.divider()
st.header("📋 Template Analysis Summary")
if st.session_state.template_analysis:
    st.info(f"**Current Template:** `{st.session_state.template_name}`")
    analysis_data = st.session_state.template_analysis
    if "error" not in analysis_data:
        # --- Display basic template summary ---
        st.markdown(f"**Filename:** `{analysis_data.get('filename', 'N/A')}`")
        # ... (display other structural summaries like page setup, TOC, etc.) ...

        # --- Display sections identified for LLM analysis ---
        if analysis_data.get("sections_with_descriptions"):
            st.subheader("📝 Sections Identified for LLM Content Analysis:")
            for sec_item in analysis_data["sections_with_descriptions"]:
                with st.expander(f"Section Title: `{sec_item['title']}`"):
                    st.markdown("**Template Description (to be used by LLM):**")
                    st.markdown(f"> {sec_item['description_from_template']}")
        elif config.get("enable_llm_validation"):
            st.caption(f"No sections with descriptions were extracted using heading style '{st.session_state.template_section_heading_style}'. LLM content validation may not be effective.")
        
        with st.expander("View Full Template Analysis (JSON)"):
            st.json(analysis_data)
    else:
        st.warning(f"Could not display template analysis: {analysis_data.get('error', 'Unknown error')}")
else:
    st.info("⬆️ Upload and analyze a template document from the sidebar to define formatting rules and content expectations.")


st.divider()
st.header("🔎 Validate User Documents")
if st.session_state.template_analysis and "error" not in st.session_state.template_analysis:
    user_doc_files = st.file_uploader(
        "2. Upload User Documents to Validate (.docx)", 
        type=["docx"], 
        accept_multiple_files=True, 
        key="user_docs_uploader"
    )

    if user_doc_files:
        for user_doc_file in user_doc_files:
            st.subheader(f"Validation Feedback for: `{user_doc_file.name}`")
            with st.spinner(f"Analyzing and validating '{user_doc_file.name}'... This may take a moment, especially if LLM is enabled."):
                user_bytes_data = user_doc_file.getvalue() # Get bytes once
                
                # Analyze the user document structure
                user_file_like_object_for_analysis = io.BytesIO(user_bytes_data)
                user_file_like_object_for_analysis.name = user_doc_file.name
                user_doc_analysis = analyze_document_structure(user_file_like_object_for_analysis, is_template=False) # Not a template

                if "error" in user_doc_analysis:
                    st.error(f"Error analyzing document '{user_doc_file.name}': {user_doc_analysis['error']}")
                    continue 

                # Store raw bytes for potential re-parsing by LLM section extractor
                if config.get("enable_llm_validation"):
                    user_doc_analysis["raw_bytes"] = user_bytes_data 
                
                # Perform validation
                feedback_messages = validate_document_against_template(
                    user_doc_analysis, 
                    st.session_state.template_analysis, 
                    config,
                    llm_model_instance_for_validation=llm_model_global_instance if config.get("enable_llm_validation") else None
                )

                # Display feedback
                for msg in feedback_messages:
                    if msg.startswith("✅") or "No major discrepancies" in msg:
                        st.success(msg, icon="🎉")
                    elif msg.startswith("🔶") or msg.startswith("ℹ️"):
                        st.info(msg)
                    elif msg.startswith("⚠️") or msg.startswith("❌") or msg.startswith("❓"):
                        st.warning(msg) # Using warning for errors too for visibility
                    elif "--- 🧠 LLM Content Relevance Analysis ---" in msg:
                        st.markdown(f"**{msg}**") # Make LLM section header bold
                    elif "└── LLM Explanation:" in msg or "└── LLM Error:" in msg:
                         st.markdown(f"&nbsp;&nbsp;&nbsp;&nbsp;{msg}") # Indent LLM explanations
                    else:
                        st.info(f"➡️ {msg}") # Default for other messages
                
                with st.expander("View Full Analysis of this User Document (JSON)"):
                    st.json(user_doc_analysis)
            st.markdown("---") # Separator
    elif st.session_state.template_analysis:
        st.info("⬆️ Upload one or more user documents to validate them against the loaded template.")
else:
    st.warning("⚠️ Please upload and analyze a template document first.")

st.sidebar.divider()
st.sidebar.markdown("Developed with AI assistance.")
st.sidebar.caption("Ensure template section descriptions are clear for best LLM results.")
